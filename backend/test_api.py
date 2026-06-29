import os
import requests
import docx

BASE_URL = "http://localhost:8000"

def create_test_docx(filename="coding_standards.docx"):
    """Generates a test Word document with paragraphs and tables."""
    print(f"Generating test DOCX: {filename}...")
    doc = docx.Document()
    doc.add_heading("ThinkStack Software Guidelines", 0)
    
    doc.add_heading("Section 1: Python Standards", level=1)
    doc.add_paragraph(
        "ThinkStack engineering standardizes on Python 3.13 for all AI development. "
        "All software must adhere to modular designs, clear variable definitions, and include functional docstrings."
    )
    
    doc.add_heading("Section 2: Frontend Tech Stack", level=1)
    doc.add_paragraph(
        "For UI development, we utilize React, TypeScript, and Vite. "
        "Styling should follow premium vanilla CSS guidelines, emphasizing glassmorphism "
        "and obsidian-dark color palettes to keep the app aesthetic consistent."
    )
    
    doc.add_heading("Section 3: Database Guidelines", level=1)
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Database Service'
    hdr_cells[1].text = 'Intended Usage'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Qdrant Local'
    row_cells[1].text = 'Vector indexing and semantic RAG search'
    
    row_cells2 = table.add_row().cells
    row_cells2[0].text = 'PostgreSQL'
    row_cells2[1].text = 'Relational metadata and history logs'
    
    doc.save(filename)
    print(f"Successfully generated {filename}")

def test_health():
    print("\n[Test] Checking /api/health...")
    res = requests.get(f"{BASE_URL}/api/health")
    print(f"Status: {res.status_code}, Response: {res.json()}")
    assert res.status_code == 200

def test_upload_pdf():
    print("\n[Test] Uploading handbook.pdf...")
    pdf_path = "handbook.pdf"
    if not os.path.exists(pdf_path):
        print(f"Error: {pdf_path} not found. Run create_test_pdf.py first.")
        return
        
    with open(pdf_path, "rb") as f:
        files = {"file": (pdf_path, f, "application/pdf")}
        res = requests.post(f"{BASE_URL}/api/upload", files=files)
    print(f"Status: {res.status_code}, Response: {res.json()}")
    assert res.status_code == 200

def test_upload_docx():
    print("\n[Test] Uploading coding_standards.docx...")
    docx_path = "coding_standards.docx"
    create_test_docx(docx_path)
    
    with open(docx_path, "rb") as f:
        files = {"file": (docx_path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = requests.post(f"{BASE_URL}/api/upload", files=files)
    print(f"Status: {res.status_code}, Response: {res.json()}")
    assert res.status_code == 200

def test_upload_url():
    print("\n[Test] Submitting webpage URL to index...")
    payload = {"url": "https://example.com"}
    res = requests.post(f"{BASE_URL}/api/url", json=payload)
    print(f"Status: {res.status_code}, Response: {res.json()}")
    assert res.status_code == 200

def test_list_documents():
    print("\n[Test] Fetching document inventory...")
    res = requests.get(f"{BASE_URL}/api/documents")
    print(f"Status: {res.status_code}")
    docs = res.json()
    for doc in docs:
        print(f"- {doc['name']} | Type: {doc['file_type'].upper()} | Blocks/Pages: {doc['total_pages']}")
    assert res.status_code == 200

def test_query_docx():
    print("\n[Test] Querying Word DOCX details: 'What is our Python style guide?'...")
    payload = {"query": "What is our Python style guide?"}
    res = requests.post(f"{BASE_URL}/api/query", json=payload)
    print(f"Status: {res.status_code}")
    data = res.json()
    print("Answer:")
    print(data.get("answer"))
    print("\nCitations:")
    for cit in data.get("citations", []):
        print(f"- Source: {cit['document_name']} | Type: {cit['file_type'].upper()} (Page {cit['page_number']}) - Score: {cit['score']:.4f}")

def test_query_url():
    print("\n[Test] Querying Web URL details: 'What is Example Domain?'...")
    payload = {"query": "What is Example Domain?"}
    res = requests.post(f"{BASE_URL}/api/query", json=payload)
    print(f"Status: {res.status_code}")
    data = res.json()
    print("Answer:")
    print(data.get("answer"))
    print("\nCitations:")
    for cit in data.get("citations", []):
         print(f"- Source: {cit['document_name']} | Type: {cit['file_type'].upper()} (URL: {cit['source_info'] or 'N/A'}) - Score: {cit['score']:.4f}")

if __name__ == "__main__":
    print("Starting ThinkStack HTTP API Phase 2 integration tests...")
    try:
        test_health()
        test_upload_pdf()
        test_upload_docx()
        test_upload_url()
        test_list_documents()
        test_query_docx()
        test_query_url()
        print("\nAll Phase 2 API tests completed successfully!")
    except Exception as e:
        print(f"\nPhase 2 Integration tests failed: {e}")
