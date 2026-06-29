import os
import re
import docx
import pypdf
import requests
from bs4 import BeautifulSoup

def load_pdf(file_path: str) -> list[dict]:
    """Extracts text page by page from a PDF file."""
    pages = []
    try:
        reader = pypdf.PdfReader(file_path)
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({
                "page_number": idx + 1,
                "text": text
            })
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")
    return pages

def load_docx(file_path: str) -> list[dict]:
    """Extracts text block-by-block from paragraphs and tables in a Word file."""
    pages = []
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Read all text paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Read all table contents
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    # De-duplicate cell values that span multiple merged columns
                    unique_cells = []
                    for cell in row_text:
                        if not unique_cells or unique_cells[-1] != cell:
                            unique_cells.append(cell)
                    full_text.append(" | ".join(unique_cells))

        # Re-group DOCX paragraphs into logical "blocks" (pages) of approx. 1200 characters
        current_block = []
        block_idx = 1
        for line in full_text:
            current_block.append(line)
            if len("\n".join(current_block)) >= 1200:
                pages.append({
                    "page_number": block_idx,
                    "text": "\n".join(current_block)
                })
                current_block = []
                block_idx += 1
                
        if current_block:
            pages.append({
                "page_number": block_idx,
                "text": "\n".join(current_block)
            })
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")
    return pages

def load_text_or_md(file_path: str) -> list[dict]:
    """Reads TXT or Markdown file content and divides it into page-like blocks."""
    pages = []
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        block_size = 1200
        # Split text into 1200-character blocks
        for idx, start in enumerate(range(0, len(text), block_size)):
            pages.append({
                "page_number": idx + 1,
                "text": text[start:start + block_size]
            })
    except Exception as e:
        raise ValueError(f"Failed to read text file: {str(e)}")
    return pages

def load_url(url: str) -> dict:
    """Crawls a website URL, extracts clean body text, and retrieves the page title."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/91.0.4472.124 Safari/537.36"
        )
    }
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Remove script tags, stylesheets, and heavy structural header/footer layout elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            element.decompose()
            
        # Get clean inner text
        raw_text = soup.get_text(separator=" ")
        
        # Clean multi-spaces and empty lines
        clean_text = re.sub(r'\s+', ' ', raw_text).strip()
        
        # Extract title
        title = "Webpage"
        if soup.title and soup.title.string:
            title = soup.title.string.strip()
            
        return {
            "title": title,
            "text": clean_text
        }
    except Exception as e:
        raise ValueError(f"Failed to crawl website URL: {str(e)}")
